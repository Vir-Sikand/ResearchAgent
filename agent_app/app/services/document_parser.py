"""
Document parsing service for multiple formats.
Supports: PDF, DOCX, PNG, JPG
All documents are processed through VILA VLM for optimal text extraction.
"""
import os
import base64
from typing import List, Dict
from io import BytesIO
from .nim_vlm_parse import parse_page_markdown_bbox


def parse_document(file_content: bytes, filename: str, file_type: str) -> Dict[str, any]:
    """
    Parse a document based on its type and extract text content using VLM.
    
    Supported formats:
    - PDF: Converts each page to image, processes with VILA
    - DOCX: Converts to PDF then processes like PDF
    - PNG/JPG: Direct VLM processing
    
    Args:
        file_content: Raw file bytes
        filename: Original filename
        file_type: MIME type or extension
        
    Returns:
        Dict with parsed_text, metadata, and parsing_method
    """
    file_lower = filename.lower()
    
    # Determine parsing method based on file type
    if file_type.startswith('application/pdf') or file_lower.endswith('.pdf'):
        return _parse_pdf(file_content, filename)
    
    elif file_type.startswith('image/') or file_lower.endswith(('.png', '.jpg', '.jpeg')):
        return _parse_image(file_content, filename)
    
    elif (file_type.startswith('application/vnd.openxmlformats-officedocument.wordprocessingml') 
          or file_lower.endswith('.docx')):
        return _parse_docx(file_content, filename)
    
    else:
        raise ValueError(f"Unsupported file type: {file_type}. Supported: PDF, DOCX, PNG, JPG")


def _parse_pdf(pdf_bytes: bytes, filename: str) -> Dict:
    """Parse PDF using VLM - converts each page to image for VLM processing."""
    try:
        import fitz
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        
        all_text = []
        dpi = 150  # Good quality for VLM
        zoom = dpi / 72.0
        
        print(f"[PDF→VLM] Processing {len(doc)} pages...")
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
            img_b64 = "data:image/png;base64," + base64.b64encode(pix.tobytes("png")).decode()
            
            # Send to VLM
            parsed = parse_page_markdown_bbox(img_b64)
            blocks = parsed.get('blocks', [])
            page_text = "\n".join([
                block.get('text', '') 
                for block in blocks
                if block.get('text') and not block.get('error')
            ])
            
            if page_text.strip():
                all_text.append(f"[Page {page_num + 1}]\n{page_text}")
            else:
                all_text.append(f"[Page {page_num + 1}]\n[VLM: No text detected]")
        
        return {
            'parsed_text': "\n\n".join(all_text),
            'metadata': {
                'filename': filename,
                'pages': len(doc),
                'size': len(pdf_bytes)
            },
            'parsing_method': 'vlm_pdf'
        }
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")


def _parse_image(img_bytes: bytes, filename: str) -> Dict:
    """Parse image using VLM (NVIDIA NeMo Retriever)."""
    try:
        print(f"[IMAGE→VLM] Processing {filename}...")
        
        # Convert to base64
        img_b64 = "data:image/png;base64," + base64.b64encode(img_bytes).decode()
        
        # Parse with VLM
        parsed = parse_page_markdown_bbox(img_b64)
        
        # Extract text from blocks
        blocks = parsed.get('blocks', [])
        text = "\n".join([
            block.get('text', '') 
            for block in blocks
            if block.get('text') and not block.get('error')
        ])
        
        if not text.strip():
            text = "[VLM: No text detected in image]"
        
        return {
            'parsed_text': text,
            'metadata': {
                'filename': filename,
                'size': len(img_bytes)
            },
            'parsing_method': 'vlm_image'
        }
    except Exception as e:
        raise ValueError(f"Failed to parse image: {str(e)}")




def _parse_docx(docx_bytes: bytes, filename: str) -> Dict:
    """
    Parse DOCX by converting to PDF first, then processing like a PDF.
    This preserves layout, formatting, and handles complex documents.
    """
    try:
        import fitz  # PyMuPDF
        from docx2pdf import convert
        import tempfile
        import os
        
        print(f"[DOCX→PDF→VLM] Converting {filename} to PDF, then processing with VLM...")
        
        # Create temporary files for DOCX and PDF
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as docx_temp:
            docx_temp.write(docx_bytes)
            docx_path = docx_temp.name
        
        pdf_path = docx_path.replace('.docx', '.pdf')
        
        try:
            # Convert DOCX to PDF
            convert(docx_path, pdf_path)
            
            # Read the PDF
            with open(pdf_path, 'rb') as pdf_file:
                pdf_bytes = pdf_file.read()
            
            # Use the PDF parser
            result = _parse_pdf(pdf_bytes, filename.replace('.docx', '.pdf'))
            result['metadata']['original_format'] = 'docx'
            result['metadata']['filename'] = filename  # Restore original filename
            result['parsing_method'] = 'vlm_docx_via_pdf'
            
            return result
            
        finally:
            # Cleanup temp files
            if os.path.exists(docx_path):
                os.unlink(docx_path)
            if os.path.exists(pdf_path):
                os.unlink(pdf_path)
                
    except ImportError as ie:
        # Fallback: If docx2pdf not available, use simpler approach
        print(f"[DOCX] docx2pdf not available, using direct text extraction: {ie}")
        return _parse_docx_simple(docx_bytes, filename)
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


def _parse_docx_simple(docx_bytes: bytes, filename: str) -> Dict:
    """
    Fallback DOCX parser - extracts text and tables directly without PDF conversion.
    """
    try:
        from docx import Document
        
        print(f"[DOCX] Simple text extraction for {filename}...")
        
        doc = Document(BytesIO(docx_bytes))
        
        all_text = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                all_text.append(para.text)
        
        # Extract tables
        if doc.tables:
            all_text.append("\n[Tables]")
            for table_idx, table in enumerate(doc.tables, 1):
                all_text.append(f"\nTable {table_idx}:")
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    all_text.append(row_text)
        
        return {
            'parsed_text': "\n\n".join(all_text) if all_text else "[No text extracted]",
            'metadata': {
                'filename': filename,
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'size': len(docx_bytes)
            },
            'parsing_method': 'docx_simple'
        }
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


def parse_multiple_documents(files: List[tuple]) -> Dict:
    """
    Parse multiple documents and combine their content.
    
    Args:
        files: List of (file_content, filename, file_type) tuples
        
    Returns:
        Combined parsed text and metadata
    """
    parsed_docs = []
    all_text = []
    
    for file_content, filename, file_type in files:
        try:
            result = parse_document(file_content, filename, file_type)
            parsed_docs.append({
                'filename': filename,
                'method': result['parsing_method'],
                'metadata': result['metadata']
            })
            all_text.append(f"[Document: {filename}]\n{result['parsed_text']}")
        except Exception as e:
            # Continue with other files even if one fails
            parsed_docs.append({
                'filename': filename,
                'error': str(e)
            })
    
    return {
        'combined_text': "\n\n" + "="*60 + "\n\n".join(all_text),
        'documents': parsed_docs,
        'total_documents': len(files),
        'successful_parses': len([d for d in parsed_docs if 'error' not in d])
    }

